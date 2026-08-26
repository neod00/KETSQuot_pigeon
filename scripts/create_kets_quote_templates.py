from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / 'web-app' / 'public' / 'templates'
NAVY = '000080'
TEAL = 'E0FFFF'


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color='000000', size='8'):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in('w:tcBorders')
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tc_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = 'w:' + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), size)
        element.set(qn('w:color'), color)


def set_cell_width(cell, width_cm):
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(int(width_cm * 567)))
    tc_w.set(qn('w:type'), 'dxa')


def set_font(run, size=9, bold=False, color=None):
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Malgun Gothic')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Malgun Gothic')
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_text(cell, text, *, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.CENTER, size=8.5):
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_pricing_table(doc, heading, prefix):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.5)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(heading)
    set_font(run, size=9.5)

    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    headers = ['구 분', '검증 일수', '검증 비용', '비 고']
    widths = [6.1, 3.1, 4.6, 3.2]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_width(cell, widths[i])
        set_cell_border(cell)
        set_cell_shading(cell, 'F8F9FA')
        add_text(cell, header, bold=True)

    rows = [
        ('신 청 비', 'N/A', '면제(720,000원)', ''),
        ('1단계(개요파악, 계획수립)', '{' + prefix + '_s1_days} Manday', '{' + prefix + '_s1_cost}원', ''),
        ('2단계(문서검토, 현장검증)', '{' + prefix + '_s2_days} Manday', '{' + prefix + '_s2_cost}원', ''),
        ('3단계(검증결과 정리/평가 등)', '{' + prefix + '_s3_days} Manday', '{' + prefix + '_s3_cost}원', ''),
        ('제경비', '-', '{' + prefix + '_expenses}원', ''),
        ('합 계', '{' + prefix + '_total_days} Manday', '{' + prefix + '_total_cost}원', 'VAT {vat_type}'),
        ('최종 제안금액', '{' + prefix + '_total_days} Manday', '{' + prefix + '_final_cost}원', 'VAT {vat_type}'),
    ]
    for index, values in enumerate(rows):
        row = table.add_row()
        if index == 5:
            fill, color, border_size = TEAL, None, '12'
        elif index == 6:
            fill, color, border_size = NAVY, 'FFFFFF', '12'
        else:
            fill, color, border_size = None, None, '8'
        for i, value in enumerate(values):
            cell = row.cells[i]
            set_cell_width(cell, widths[i])
            set_cell_border(cell, size=border_size)
            if fill:
                set_cell_shading(cell, fill)
            add_text(cell, value, bold=index >= 5, color=color,
                     align=WD_ALIGN_PARAGRAPH.RIGHT if i == 2 else WD_ALIGN_PARAGRAPH.CENTER)


def add_footer(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run('LRQA Korea  |  02-3703-7514  |  www.lrqa.com')
    set_font(run, size=8, color='666666')


def create_template(kind, filename):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    add_footer(section)

    top_line = doc.add_paragraph()
    top_line.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = top_line.add_run('LRQA')
    set_font(run, size=10, bold=True, color='008C95')

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(14)
    run = title.add_run('{document_title}')
    set_font(run, size=20, bold=True)

    info = doc.add_table(rows=3, cols=4)
    info.autofit = False
    info_values = [
        ('수신:', '{company_name} 귀하', '문서번호:', '{proposal_no}'),
        ('참조:', '{contact_person}', '발행일:', '{proposal_date}'),
        ('발신:', '로이드인증원(LRQA)', '', ''),
    ]
    for r, values in enumerate(info_values):
        for c, value in enumerate(values):
            cell = info.rows[r].cells[c]
            set_cell_width(cell, [1.1, 6.4, 1.4, 7.0][c])
            add_text(cell, value,
                     align=WD_ALIGN_PARAGRAPH.RIGHT if c in (2, 3) else WD_ALIGN_PARAGRAPH.LEFT,
                     size=9)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(8)
    ppr = rule._p.get_or_add_pPr()
    borders = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:color'), '000000')
    borders.append(bottom)
    ppr.append(borders)

    greeting = doc.add_paragraph()
    greeting.paragraph_format.space_before = Pt(10)
    greeting.paragraph_format.space_after = Pt(5)
    run = greeting.add_run('1. 귀 사의 지속가능한 발전을 기원합니다.')
    set_font(run, size=9.5)
    detail = doc.add_paragraph()
    detail.paragraph_format.space_after = Pt(8)
    run = detail.add_run('2. 귀 사에 대한 {scope_text} 검증과 관련하여, 예상되는 검증 심사 내역을 아래와 같이 송부드리오니 업무 참조하시기 바랍니다.')
    set_font(run, size=9.5)

    scope = doc.add_paragraph()
    scope.paragraph_format.space_before = Pt(2)
    scope.paragraph_format.space_after = Pt(2)
    run = scope.add_run('(1) 검증 범위: {scope_text}')
    set_font(run, size=9)
    standard = doc.add_paragraph()
    standard.paragraph_format.space_after = Pt(4)
    run = standard.add_run('(2) 심사 기준: 온실가스 배출권거래제의 배출량 보고 및 인증에 관한 지침')
    set_font(run, size=9)

    if kind in ('statement', 'combined'):
        add_pricing_table(doc, '1) 온실가스 명세서', 'statement')
    if kind in ('plan', 'combined'):
        number = '2)' if kind == 'combined' else '1)'
        add_pricing_table(doc, number + ' 배출량산정계획서', 'plan')

    if kind == 'combined':
        total = doc.add_table(rows=1, cols=2)
        total.autofit = False
        for i, value in enumerate(['전체 최종 제안금액 (명세서 + 계획서)', '{combined_final_cost}원 (VAT {vat_type})']):
            cell = total.rows[0].cells[i]
            set_cell_width(cell, 9.0 if i == 0 else 8.0)
            set_cell_border(cell, color=NAVY, size='18')
            set_cell_shading(cell, NAVY)
            add_text(cell, value, bold=True, color='FFFFFF', size=9.5,
                     align=WD_ALIGN_PARAGRAPH.RIGHT if i else WD_ALIGN_PARAGRAPH.CENTER)

    notes = [
        '(5) 기타',
        '1) 심사 요율은 {audit_rate}원/ Manday 이며 상기 금액은 부가가치세(VAT)가 {vat_description} 금액입니다.',
        '2) 교통비, 숙박비, 심사원 일비 등의 제경비는 상기 제안금액에 포함되어 있습니다.',
        '3) 상기 검증비용은 업체의 상황에 따라 상호 협의 하에 조정될 수 있습니다.',
        '4) 제안서 유효기간은 제안 발행일로부터 30일 이내 입니다.',
        '{contract_note}',
        '{contact_note}',
    ]
    for index, text in enumerate(notes):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.45 if index else 0)
        paragraph.paragraph_format.space_before = Pt(2 if index else 9)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        set_font(run, size=8.8, bold=index == 0)

    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    closing.paragraph_format.space_before = Pt(12)
    run = closing.add_run('감사합니다.\n로이드인증원')
    set_font(run, size=9.5, bold=True)

    doc.save(OUTPUT / filename)


if __name__ == '__main__':
    OUTPUT.mkdir(parents=True, exist_ok=True)
    create_template('statement', 'LRQA_KETS_Statement_Quote_Template.docx')
    create_template('plan', 'LRQA_KETS_Plan_Quote_Template.docx')
    create_template('combined', 'LRQA_KETS_Combined_Quote_Template.docx')
